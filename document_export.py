"""
Builds downloadable .docx files from generated CV and cover letter text.

Uses python-docx (already a dependency for reading uploaded resumes) so no
new libraries are needed to also write documents.
"""
import io
from docx import Document
from docx.shared import Pt


def _add_paragraphs(doc: Document, text: str) -> None:
    """
    Split text into paragraphs on blank lines and add each as its own
    paragraph. Keeps a plain, readable layout rather than guessing at
    markdown-style formatting the model may or may not have used.
    """
    for block in text.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)


def build_cv_docx(name: str, cv_text: str) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading(name or "Curriculum Vitae", level=0)
    title.runs[0].font.size = Pt(20)

    _add_paragraphs(doc, cv_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def build_cover_letter_docx(name: str, cover_letter_text: str) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading("Cover Letter", level=0)
    title.runs[0].font.size = Pt(18)

    _add_paragraphs(doc, cover_letter_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
